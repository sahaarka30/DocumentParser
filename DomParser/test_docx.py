import pandas as pd
import json
from docx import *
from io import *
import zipfile
import xml.etree.ElementTree
from docx import Document
from StringIO import StringIO
import officedissector
import json
from xmltodict import parse
from xml2json import *
import xmltodict
from docx.shared import Inches
from docx.shared import *

# print(doc)
# WORD_NAMESPACE = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
# PARA = WORD_NAMESPACE + 'p'
# TEXT = WORD_NAMESPACE + 't'
# TABLE = WORD_NAMESPACE + 'tbl'
# ROW = WORD_NAMESPACE + 'tr'
# CELL = WORD_NAMESPACE + 'tc'
# with zipfile.ZipFile('easychair.docx') as docx:
#     tree = xml.etree.ElementTree.XML(docx.read('word/document.xml'))

# for text in tree.iter(PARA):
#     for r in text.iter(TEXT):
#         print (r.)


# for table in tree.iter(TABLE):
#     for row in table.iter(ROW):
#         for cell in row.iter(CELL):
#             print (''.join(node.text for node in cell.iter(TEXT)))
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

doc= Document("test1.docx")


# doc=doc.paragraphs
# print(doc)

# for p in doc.paragraphs:
    # ri=p.style.name
    # print p,":",ri
    # print p.text,p.alignment

    # try:
    #     tab = p.paragraph_format.space_before.pt
    #     tab1 = p.paragraph_format.space_after.pt
    #     print tab, ":", tab1
    # except:
    #     tab = p.paragraph_format.space_before
    #     tab1 = p.paragraph_format.space_after
    #     print tab, ":", tab1


    # for r in p.runs:
    #
    #     print r.text," : ",r.font.name," : ",r.font.cs_bold," : ",r.font.cs_italic," : ",r.font.size
        # if(f.style):
        #     f="bold"
        # elif (f.italic):
        #     f="italic"
        # else:
        #     f="None"
        # print(r.text,"::",r.style.name,"::",f)
        # print("\n")
# document.save('new-file-name.docx')
# print(doc)

# doc=doc.paragraphs
# for p in doc:
#     for r in p.runs:
#         set=iter_block_items(p)
#         print(p.text,"::",set)
#         print("\n")
# with open('easychair.docx', 'rb') as f:
#     source_stream = StringIO(str(f.read()))
# doc=Document(source_stream)

def get_detail_text_data(path):

    document = Document(path)
    content = list()
    for paragraph in document.paragraphs:
        para = dict()
        para['indent'] = paragraph.alignment
        para['first_line_indent'] = paragraph.paragraph_format.first_line_indent
        para['style_name'] = paragraph.style.name

        try:
            para['space_before'] = paragraph.paragraph_format.space_before.pt
            para['space_after'] = paragraph.paragraph_format.space_after.pt

        except:
            para['space_before'] = paragraph.paragraph_format.space_before
            para['space_after'] = paragraph.paragraph_format.space_after

        para['runs'] = list()
        for run in paragraph.runs:
            di = dict()
            di['text'] = run.text

            di['style_name'] = run.font.name
            di['style_font_name'] = run.font.name
            try:
                di['style_font_size'] = run.font.size.pt
            except:
                pass

            di['style_font_italic'] = "italic" if (run.font.cs_italic) else "None"
            di['style_font_bold'] = "bold" if (run.font.cs_bold) else "None"
            di['style_font_underline'] = "underline" if (run.font.underline) else "None"
            para['runs'].append(di)
        content.append(para)

    return content

def get_styles(data):
    styles = set()
    for paragraph in data:
        style = paragraph.get('style_name')
        styles.add(style)

    archive = zipfile.ZipFile('test1.docx', 'r')
    data = archive.read('word/styles.xml')

    data = json.dumps(xmltodict.parse(data))
    data = json.loads(data)

    # print data
    data = data['w:styles']['w:style']
    # print data

    # print json.dumps(tab_dec)
    t_count=list()
    tab_dec = archive.read('word/document.xml')
    tab_dec = json.dumps(xmltodict.parse(tab_dec))
    tab_dec = json.loads(tab_dec)
    tab_dec = tab_dec['w:document']['w:body']['w:p']
    for t in tab_dec:
        try:
            if bool(t['w:r']['w:tab']):
               t_count.append({len(t['w:r']['w:tab']),str(t['w:r']['w:t'])})
            else:
                t_count.append({1,str(t['w:r']['w:t'])})
        except:
            pass

    t_count=list(t_count[0])
    # print t_count
    for p in doc.paragraphs:
        for r in p.runs:
            print ":/",r.text
            if r.text in t_count:
                print "True"

    styles_lst = list()
    for style in styles:
        style_id = style.replace(" ", "")
        link_id = ''
        for item in data:
            if style_id == item.get('@w:styleId'):
                if item.get('w:link'):
                    link_id = item.get('w:link').get('@w:val')
                    break

        di = dict()
        di['styles'] = dict()
        di['style_name'] = style
        for item in data:
            if link_id == item.get('@w:styleId'):

                font = item.get('w:rPr')
                if font.get('w:spacing'):
                    di['styles']['spacing'] = font.get('w:spacing').get('@w:val')
                di['styles']['size'] = font.get('w:sz').get('@w:val')
                di['styles']['fonts'] = {
                    'ascii': font.get('w:rFonts').get('@w:ascii'),
                    'eastAsiaTheme': font.get('w:rFonts').get('@w:eastAsiaTheme'),
                    'hAnsi': font.get('w:rFonts').get('@w:hAnsi'),
                    'cstheme': font.get('w:rFonts').get('@w:cstheme'),
                }
                styles_lst.append(di)

                break

    return styles_lst



d=get_detail_text_data("test1.docx")
sty=get_styles(d)
# print sty

