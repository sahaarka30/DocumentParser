from docx import Document
from StringIO import StringIO

import json


with open('easychair.docx', 'rb') as f:
    source_stream = StringIO(f.read())

document = Document(source_stream)

content = list()
for paragraph in document.paragraphs:
    para = dict()
    para['left_indent'] = paragraph.paragraph_format.left_indent
    para['right_indent'] = paragraph.paragraph_format.right_indent
    para['first_line_indent'] = paragraph.paragraph_format.first_line_indent
    para['runs'] = list()
    for run in paragraph.runs:
        di = dict()
        di['text'] = run.text
        di['style_name'] = run.style.name
        di['style_font_name'] = run.style.font.name
        di['style_font_size'] = run.style.font.size
        di['style_font_italic'] = run.style.font.italic
        di['style_font_bold'] = run.style.font.bold
        di['style_font_underline'] = run.style.font.underline
        para['runs'].append(di)

    content.append(para)


# print "creating file ..."
with open('content.json', 'w') as outfile:
    json.dump(content, outfile, indent=4)
# print "file created ..."

# print "length :", len(document.paragraphs)
