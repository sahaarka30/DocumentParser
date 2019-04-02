from docx import Document
from StringIO import StringIO
import officedissector
import json
from xmltodict import parse
import zipfile
import xmltodict
from xml2json import *
from docx.shared import *

class Docx:

    def __init__(self, filename):
        self.file = filename

    def get_detail_text_data(self):
        with open(self.file, 'rb') as f:
            source_stream = StringIO(f.read())

        document = Document(source_stream)
        content = list()
        for paragraph in document.paragraphs:
            para = dict()
            para['indent'] = paragraph.alignment
            para['first_line_indent'] = paragraph.paragraph_format.first_line_indent
            para['style_name'] = paragraph.style.name

            try:
                para['space_before'] = paragraph.paragraph_format.space_before.pt
                para['space_after'] = paragraph.paragraph_format.space_after.pt

            except:
                para['space_before'] = paragraph.paragraph_format.space_before
                para['space_after'] = paragraph.paragraph_format.space_after

            para['runs'] = list()
            for run in paragraph.runs:
                di = dict()
                di['text'] = run.text
                di['text'] = di['text'].replace("\t","&emsp;")
                di['style_name'] = run.font.name
                di['style_font_name'] = run.font.name
                try:
                    di['style_font_size'] = run.font.size.pt
                except:
                    pass

                di['style_font_italic'] = "italic"if(run.font.cs_italic)else"None"
                di['style_font_bold'] = "bold"if(run.font.cs_bold)else"None"
                di['style_font_underline'] = "underline"if(run.font.underline)else"None"
                para['runs'].append(di)
            content.append(para)

        return content

    def get_images(self):
        doc = officedissector.doc.Document(self.file)
        mp = doc.main_part()

        doc_data = doc.to_json()

        parts = list()
        for document in json.loads(doc_data)['document']:
            if "parts" in document:
                parts = document['parts']
        di = dict()
        for item in parts:
            if "image" in item['content-type']:
                uri = item['uri']
                relationship = item['relationships_in'][0].split()[1].replace('[', '').replace(']', '')
                di[relationship] = uri

        return di

    def get_image_position_in_text(self):
        archive = zipfile.ZipFile(self.file, 'r')
        data = archive.read('word/document.xml')

        data  = json.dumps(xmltodict.parse(data))
        data = json.loads(data)
        data = data['w:document']['w:body']['w:p']

        content = list()
        for index, paragraph in enumerate(data):
            para_content = list()
            para_index = index
            if paragraph.get("w:r"):
                for item in paragraph.get("w:r"):
                    di = dict()
                    if type(item) is dict:
                        if  type(item.get('w:t')) is dict:
                            text = item.get('w:t').get('#text')
                        else:
                            text = item.get('w:t')
                        di['text'] = text
                    para_content.append(di)
            if "a:blip" in str(paragraph):
                di = dict()
                embed = str(paragraph).split('@r:embed')[1].split('}')[0].split()[1].replace("u'", "").replace("'", "").replace(",", "")
                di['embed'] = embed
                para_content.append(di)
            di = dict()
            di['para_content'] = para_content
            di['para_index'] = para_index
            content.append(di)

        return content

    def create_final_data(self, content, images, content_with_images):
        final_content = list()
        for content_index, content_item in enumerate(content):
            di = dict()
            di['first_line_indent'] = content_item.get('first_line_indent')
            di['indent'] = str(content_item.get('indent'))
            di['indent'] =di['indent'][0:len(di['indent'])-4]
            di['style_name'] = content_item.get('style_name')
            di['space_before']= content_item.get('space_before')
            di['space_after'] = content_item.get('space_after')
            di['data'] = list()
            content_with_images_item = content_with_images[content_index].get('para_content')
            for index, item in enumerate(content_item.get('runs')):
                data_di = dict()
                data_di['style_font_name'] = item.get('style_font_name')
                data_di['style_font_size'] = item.get('style_font_size')
                data_di['text'] = item.get('text')
                data_di['style_font_underline'] = item.get('style_font_underline')
                data_di['style_font_italic'] = item.get('style_font_italic')
                data_di['style_name'] = item.get('style_name')
                data_di['style_font_bold'] = item.get('style_font_bold')
                di['data'].append(data_di)
            if (len(content_item.get('runs')) < len(content_with_images_item)) and ('embed' in content_with_images_item[-1]):
                data_di = dict()
                data_di['image'] = images.get(content_with_images_item[-1].get('embed'))
                di['data'].append(data_di)
            final_content.append(di)

        return final_content

    def get_styles(self, data,filename):
        styles = set()
        for paragraph in data:
            style = paragraph.get('style_name')
            styles.add(style)

        archive = zipfile.ZipFile(filename, 'r')
        data = archive.read('word/styles.xml')
        data = json.dumps(xmltodict.parse(data))
        data = json.loads(data)
        data = data['w:styles']['w:style']




        styles_lst = list()
        for style in styles:
            style_id = style.replace(" ", "")
            link_id = ''
            for item in data:
                if style_id == item.get('@w:styleId'):
                    if item.get('w:link'):
                        link_id = item.get('w:link').get('@w:val')
                        break

            di = dict()
            di['styles'] = dict()
            di['style_name'] = style
            for item in data:
                if link_id == item.get('@w:styleId'):

                    font = item.get('w:rPr')
                    if font.get('w:spacing'):
                        di['styles']['spacing'] = font.get('w:spacing').get('@w:val')
                    di['styles']['size'] = font.get('w:sz').get('@w:val')
                    di['styles']['fonts'] = {
                        'ascii': font.get('w:rFonts').get('@w:ascii'),
                        'eastAsiaTheme': font.get('w:rFonts').get('@w:eastAsiaTheme'),
                        'hAnsi': font.get('w:rFonts').get('@w:hAnsi'),
                        'cstheme': font.get('w:rFonts').get('@w:cstheme'),
                    }
                    styles_lst.append(di)

                    break

        return styles_lst

    def create_content_with_styles(self, content, styles):
        final_content = list()
        for item in content:
            style_name = item.get('style_name')
            for style in styles:
                if style_name == style.get('style_name'):
                    styl = style.get('styles')
                    item['styles'] = styl
                    break

            final_content.append(item)

        return final_content



def print_file(filename, content):
    # print "creating file ..."
    with open(filename, 'w') as outfile:
        json.dump(content, outfile, indent=4)
    # print "file created ..."

if __name__ == '__main__':
    file='test1.docx'
    docx_obj = Docx(file)
    # docx_obj = Docx('./demo.docx')
    content = docx_obj.get_detail_text_data()

    images = docx_obj.get_images()
    content_with_images = docx_obj.get_image_position_in_text()
    content = docx_obj.create_final_data(content, images, content_with_images)
    styles = docx_obj.get_styles(content,file)
    # print styles
    final_content = docx_obj.create_content_with_styles(content, styles)
    # print final_content
    print_file('final_content.json', final_content)
    print_file('demo.json', content)